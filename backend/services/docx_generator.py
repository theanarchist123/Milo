import os
import io
import re
from docx import Document
from docx.shared import Pt, RGBColor
import uuid

def generate_docx_bytes(title: str, markdown_text: str) -> io.BytesIO:
    """
    Creates a beautifully formatted DOCX file from a markdown string
    and returns it as an in-memory BytesIO object. (Serverless friendly)
    """
    doc = Document()
    
    # Add Title
    title_par = doc.add_heading(title, level=0)
    
    lines = markdown_text.split('\n')
    in_code_block = False
    
    for line in lines:
        stripped_line = line.strip()
        
        # Handle code blocks toggle
        if stripped_line.startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                # Add a little spacing before code blocks
                doc.add_paragraph()
            continue
            
        if in_code_block:
            # Code block styling
            p = doc.add_paragraph()
            # Try to style it to look like code (Courier New)
            run = p.add_run(line) # preserve indentation
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # Reduce spacing inside code blocks
            p.paragraph_format.space_after = Pt(2)
            continue
            
        # Ignore empty lines
        if not stripped_line:
            doc.add_paragraph()
            continue
            
        # Headings
        if stripped_line.startswith('#'):
            level = len(stripped_line) - len(stripped_line.lstrip('#'))
            clean_text = stripped_line.lstrip('#').strip()
            # Strip bold formatting naturally if Gemini included it in the heading
            clean_text = clean_text.replace('**', '').replace('*', '') 
            doc.add_heading(clean_text, level=min(level, 4))
            continue
            
        # Unordered Lists
        if stripped_line.startswith('- ') or stripped_line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, stripped_line[2:].strip())
            continue
            
        # Ordered Lists
        if re.match(r'^\d+\.\s', stripped_line):
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, re.sub(r'^\d+\.\s', '', stripped_line).strip())
            continue
            
        # Handle bold lines that are just labels like "Aim:" without making them headings
        if stripped_line.startswith('**') and ':' in stripped_line and stripped_line.index(':') < 30:
             p = doc.add_paragraph()
             _add_formatted_text(p, stripped_line)
             continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_formatted_text(p, stripped_line)
        
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f

def create_docx_from_markdown(title: str, markdown_text: str) -> str:
    """
    Legacy wrapper for backwards compatibility with existing database rows.
    In Vercel, we don't save to disk anymore.
    """
    return f"outputs/dynamic_{uuid.uuid4().hex}.docx"

def _add_formatted_text(paragraph, text: str):
    """
    Helper to add text with basic **bold**, *italic*, and `code` formatting to a paragraph.
    It breaks down the text progressively.
    """
    # Progressive Regex splitting to handle multiple formats gracefully.
    # 1. Split by backticks for inline code
    code_parts = re.split(r'(`[^`]+`)', text)
    
    for c_part in code_parts:
        if c_part.startswith('`') and c_part.endswith('`'):
            run = paragraph.add_run(c_part[1:-1])
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(0xCC, 0x44, 0x44)
        else:
            # 2. Split by bold
            bold_parts = re.split(r'(\*\*.*?\*\*)', c_part)
            for b_part in bold_parts:
                if b_part.startswith('**') and b_part.endswith('**'):
                    run = paragraph.add_run(b_part[2:-2])
                    run.bold = True
                else:
                    # 3. Split by italic
                    italic_parts = re.split(r'(\*[^\*]+\*)', b_part)
                    for i_part in italic_parts:
                        if i_part.startswith('*') and i_part.endswith('*'):
                            # Ensure we didn't just match an empty string
                            if len(i_part) > 2:
                                run = paragraph.add_run(i_part[1:-1])
                                run.italic = True
                            else:
                                paragraph.add_run('*')
                        else:
                            # Final text dump
                            
                            # Replace any remaining markdown artifacts if present safely
                            # (Like unparsed single underscores if they exist, but we leave them as is for now)
                            paragraph.add_run(i_part)
